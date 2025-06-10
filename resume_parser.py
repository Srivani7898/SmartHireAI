import fitz  # fitz
import docx
import re
import spacy
from typing import Dict, List, Optional
import os

class ResumeParser:
    def __init__(self):
        """Initialize the resume parser with spaCy model."""
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except OSError:
            # Fallback if spaCy model is not installed
            self.nlp = None
            print("Warning: spaCy model 'en_core_web_sm' not found. Some features may be limited.")
        
        # Define skill patterns and keywords
        self.skill_patterns = [
            # Programming languages
            r'\b(?:python|java|javascript|c\+\+|c#|php|ruby|go|rust|swift|kotlin|scala|r|matlab)\b',
            # Web technologies
            r'\b(?:html|css|react|angular|vue|node\.?js|express|django|flask|spring|laravel)\b',
            # Databases
            r'\b(?:mysql|postgresql|mongodb|redis|elasticsearch|oracle|sql\s*server|sqlite)\b',
            # Cloud & DevOps
            r'\b(?:aws|azure|gcp|docker|kubernetes|jenkins|git|github|gitlab|terraform)\b',
            # Data Science & ML
            r'\b(?:pandas|numpy|scikit-learn|tensorflow|pytorch|keras|spark|hadoop|tableau|power\s*bi)\b',
            # Other technologies
            r'\b(?:linux|windows|macos|agile|scrum|jira|confluence|slack|microsoft\s*office)\b'
        ]
        
        self.education_keywords = [
            'bachelor', 'master', 'phd', 'doctorate', 'degree', 'diploma', 'certificate',
            'b.tech', 'b.e', 'm.tech', 'm.e', 'mba', 'mca', 'bca', 'b.sc', 'm.sc',
            'engineering', 'computer science', 'information technology', 'software',
            'university', 'college', 'institute', 'school'
        ]
        
        self.experience_keywords = [
            'experience', 'worked', 'developed', 'managed', 'led', 'created', 'designed',
            'implemented', 'built', 'maintained', 'optimized', 'collaborated', 'achieved',
            'years', 'months', 'intern', 'internship', 'trainee', 'junior', 'senior',
            'lead', 'manager', 'developer', 'engineer', 'analyst', 'consultant'
        ]
    
    def parse_resume(self, file_path: str) -> Optional[str]:
        """
        Parse resume from PDF or DOCX file and extract text.
        
        Args:
            file_path (str): Path to the resume file
            
        Returns:
            Optional[str]: Extracted text or None if parsing fails
        """
        try:
            file_extension = os.path.splitext(file_path)[1].lower()
            
            if file_extension == '.pdf':
                return self._parse_pdf(file_path)
            elif file_extension == '.docx':
                return self._parse_docx(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            print(f"Error parsing resume {file_path}: {str(e)}")
            return None
    
    def _parse_pdf(self, file_path: str) -> str:
        """Extract text from PDF file using PyMuPDF."""
        text = ""
        try:
            doc = PyMuPDF.open(file_path)
            for page in doc:
                text += page.get_text()
            doc.close()
        except Exception as e:
            raise Exception(f"Error parsing PDF: {str(e)}")
        
        return text.strip()
    
    def _parse_docx(self, file_path: str) -> str:
        """Extract text from DOCX file using python-docx."""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
            
        except Exception as e:
            raise Exception(f"Error parsing DOCX: {str(e)}")
    
    def extract_information(self, resume_text: str) -> Dict[str, List[str]]:
        """
        Extract structured information from resume text.
        
        Args:
            resume_text (str): Raw resume text
            
        Returns:
            Dict[str, List[str]]: Dictionary containing extracted information
        """
        resume_text_lower = resume_text.lower()
        
        # Extract skills
        skills = self._extract_skills(resume_text_lower)
        
        # Extract education information
        education = self._extract_education(resume_text_lower)
        
        # Extract experience information
        experience = self._extract_experience(resume_text_lower)
        
        # Extract contact information
        contact = self._extract_contact_info(resume_text)
        
        # Extract years of experience
        years_exp = self._extract_years_of_experience(resume_text_lower)
        
        return {
            'skills': skills,
            'education': education,
            'experience': experience,
            'contact': contact,
            'years_of_experience': years_exp
        }
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extract technical skills from resume text."""
        skills = set()
        
        # Use regex patterns to find skills
        for pattern in self.skill_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            skills.update([match.lower() for match in matches])
        
        # Use spaCy for named entity recognition if available
        if self.nlp:
            doc = self.nlp(text)
            for ent in doc.ents:
                if ent.label_ in ['ORG', 'PRODUCT']:  # Organizations and products might be technologies
                    skills.add(ent.text.lower())
        
        # Clean and filter skills
        cleaned_skills = []
        for skill in skills:
            skill = skill.strip()
            if len(skill) > 1 and skill.isalpha() or '.' in skill or '+' in skill or '#' in skill:
                cleaned_skills.append(skill)
        
        return list(set(cleaned_skills))[:20]  # Return top 20 unique skills
    
    def _extract_education(self, text: str) -> List[str]:
        """Extract education information from resume text."""
        education_info = []
        
        # Look for education keywords
        for keyword in self.education_keywords:
            if keyword in text:
                # Find sentences containing education keywords
                sentences = text.split('.')
                for sentence in sentences:
                    if keyword in sentence:
                        education_info.append(sentence.strip())
        
        # Extract degree patterns
        degree_patterns = [
            r'b\.?tech|bachelor of technology',
            r'm\.?tech|master of technology',
            r'b\.?e\.?|bachelor of engineering',
            r'm\.?e\.?|master of engineering',
            r'mba|master of business administration',
            r'mca|master of computer applications',
            r'bca|bachelor of computer applications',
            r'b\.?sc|bachelor of science',
            r'm\.?sc|master of science',
            r'phd|doctorate'
        ]
        
        for pattern in degree_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            education_info.extend(matches)
        
        return list(set(education_info))[:10]  # Return unique education info
    
    def _extract_experience(self, text: str) -> List[str]:
        """Extract work experience information from resume text."""
        experience_info = []
        
        # Look for experience keywords
        for keyword in self.experience_keywords:
            if keyword in text:
                # Find sentences containing experience keywords
                sentences = text.split('.')
                for sentence in sentences:
                    if keyword in sentence and len(sentence.strip()) > 10:
                        experience_info.append(sentence.strip())
        
        # Extract job titles and companies
        job_patterns = [
            r'(?:software|web|mobile|data|ml|ai)\s+(?:engineer|developer|analyst|scientist)',
            r'(?:senior|junior|lead|principal)\s+(?:engineer|developer|analyst)',
            r'(?:project|product|technical)\s+(?:manager|lead)',
            r'(?:full\s*stack|front\s*end|back\s*end)\s+developer',
            r'(?:data|business|system)\s+analyst'
        ]
        
        for pattern in job_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            experience_info.extend(matches)
        
        return list(set(experience_info))[:15]  # Return unique experience info
    
    def _extract_contact_info(self, text: str) -> Dict[str, str]:
        """Extract contact information from resume text."""
        contact_info = {}
        
        # Email pattern
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            contact_info['email'] = emails[0]
        
        # Phone pattern
        phone_pattern = r'(?:\+?1[-.\s]?)?$$?[0-9]{3}$$?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}'
        phones = re.findall(phone_pattern, text)
        if phones:
            contact_info['phone'] = phones[0]
        
        # LinkedIn pattern
        linkedin_pattern = r'linkedin\.com/in/[\w-]+'
        linkedin = re.findall(linkedin_pattern, text, re.IGNORECASE)
        if linkedin:
            contact_info['linkedin'] = linkedin[0]
        
        # GitHub pattern
        github_pattern = r'github\.com/[\w-]+'
        github = re.findall(github_pattern, text, re.IGNORECASE)
        if github:
            contact_info['github'] = github[0]
        
        return contact_info
    
    def _extract_years_of_experience(self, text: str) -> Optional[int]:
        """Extract years of experience from resume text."""
        # Pattern to match years of experience
        exp_patterns = [
            r'(\d+)\s*(?:\+)?\s*years?\s+(?:of\s+)?experience',
            r'experience\s+(?:of\s+)?(\d+)\s*(?:\+)?\s*years?',
            r'(\d+)\s*(?:\+)?\s*yrs?\s+(?:of\s+)?(?:exp|experience)',
            r'(?:exp|experience)\s+(?:of\s+)?(\d+)\s*(?:\+)?\s*yrs?'
        ]
        
        for pattern in exp_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                try:
                    return int(matches[0])
                except ValueError:
                    continue
        
        return None
    
    def get_resume_summary(self, resume_text: str) -> Dict[str, any]:
        """
        Get a comprehensive summary of the resume.
        
        Args:
            resume_text (str): Raw resume text
            
        Returns:
            Dict[str, any]: Summary information
        """
        info = self.extract_information(resume_text)
        
        return {
            'word_count': len(resume_text.split()),
            'character_count': len(resume_text),
            'skills_count': len(info['skills']),
            'education_count': len(info['education']),
            'experience_count': len(info['experience']),
            'has_contact_info': bool(info['contact']),
            'years_of_experience': info['years_of_experience'],
            'top_skills': info['skills'][:5],
            'extracted_info': info
        }
